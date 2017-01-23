

# Examples from https://automatetheboringstuff.com/chapter13/

"""
#    >>> import docx
# ❶ >>> doc = docx.Document('demo.docx')
# ❷ >>> len(doc.paragraphs)
#    7
# ❸ >>> doc.paragraphs[0].text
#    'Document Title'
# ❹ >>> doc.paragraphs[1].text
#    'A plain paragraph with some bold and some italic'
# ❺ >>> len(doc.paragraphs[1].runs)
#    4
# ❻ >>> doc.paragraphs[1].runs[0].text
#    'A plain paragraph with some '
# ❼ >>> doc.paragraphs[1].runs[1].text
#    'bold'
# ❽ >>> doc.paragraphs[1].runs[2].text
#    ' and some '
# ➒ >>> doc.paragraphs[1].runs[3].text
#    'italic'
"""

import docx

